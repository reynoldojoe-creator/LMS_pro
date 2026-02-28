from fastapi import APIRouter, Form, HTTPException, Depends, Response
from typing import Optional, List, Dict, Any
import json
import uuid
import io
from docx import Document
from datetime import datetime
from sqlalchemy.orm import Session
from ...models import database
from ...services.rubric_service import rubric_service

router = APIRouter(prefix="/rubrics", tags=["rubrics"])

def get_db():
    db = database.SessionLocal()
    try:
        yield db
    finally:
        db.close()

@router.post("")
async def create_rubric(
    request: Dict[str, Any],
    db: Session = Depends(get_db)
):
    """
    Create rubric — accepts plain JSON dict.
    Handles both camelCase (from some frontends) and snake_case keys.
    """
    try:
        # Helper to get value from either camelCase or snake_case key
        def get(camel: str, snake: str, default=None):
            return request.get(camel, request.get(snake, default))

        rubric_id = str(uuid.uuid4())
        title = get("title", "name", "Untitled")
        subject_id = int(get("subjectId", "subject_id", 0))
        exam_type = get("examType", "exam_type", "final")
        total_marks = int(get("totalMarks", "total_marks", 100))
        duration = int(get("duration", "duration_minutes", 180))

        # Serialize dict/list fields to JSON strings for Text columns
        def to_json(val, default="{}"):
            if val is None:
                return default
            if isinstance(val, (dict, list)):
                return json.dumps(val)
            if isinstance(val, str):
                return val
            return json.dumps(val)

        sections = to_json(
            get("sections", "question_distribution", {}),
            "{}"
        )
        co_dist = to_json(get("coDistribution", "co_distribution", None), None)
        co_req = to_json(get("coRequirements", "co_requirements", None), None)
        lo_dist = to_json(get("loDistribution", "lo_distribution", None), None)
        diff_dist = to_json(get("difficultyDistribution", "difficulty_distribution", {}))
        bloom_dist = to_json(get("bloomDistribution", "bloom_distribution", None), None)
        units = to_json(get("unitsCovered", "units_covered", []), "[]")
        assign_cfg = to_json(get("assignmentConfig", "assignment_config", None), None)

        rubric = database.Rubric(
            id=rubric_id,
            title=title,
            name=title,  # Legacy column, NOT NULL in physical DB
            subject_id=subject_id,
            exam_type=exam_type,
            total_marks=total_marks,
            duration=duration,
            sections=sections,
            difficulty_distribution=diff_dist,
            co_requirements=co_req,
            co_distribution=co_dist,
            lo_distribution=lo_dist,
            bloom_distribution=bloom_dist,
            units_covered=units,
            assignment_config=assign_cfg,
            status="created",
            created_at=datetime.utcnow(),
        )
        db.add(rubric)
        db.commit()
        db.refresh(rubric)

        # Create Rubric Items from sections
        if sections and sections != "{}":
            sections_data = json.loads(sections)
            for q_type, config in sections_data.items():
                if isinstance(config, dict):
                    item = database.RubricItem(
                        rubric_id=rubric.id,
                        question_type=q_type,
                        marks=config.get("marks_each", 1),
                        count=config.get("count", 0),
                        difficulty=config.get("difficulty", "medium")
                    )
                    db.add(item)
            db.commit()

        return {
            "id": str(rubric.id),
            "title": rubric.title,
            "name": rubric.name,
            "subjectId": str(rubric.subject_id),
            "examType": rubric.exam_type,
            "duration": rubric.duration,
            "totalMarks": rubric.total_marks,
            "sections": json.loads(rubric.sections) if rubric.sections else {},
            "difficultyDistribution": json.loads(rubric.difficulty_distribution) if rubric.difficulty_distribution else {},
            "coRequirements": json.loads(rubric.co_requirements) if rubric.co_requirements else [],
            "coDistribution": json.loads(rubric.co_distribution) if rubric.co_distribution else {},
            "loDistribution": json.loads(rubric.lo_distribution) if rubric.lo_distribution else {},
            "bloomDistribution": json.loads(rubric.bloom_distribution) if rubric.bloom_distribution else {},
            "unitsCovered": json.loads(rubric.units_covered) if rubric.units_covered else [],
            "status": rubric.status,
            "createdAt": rubric.created_at.isoformat() if rubric.created_at else None,
        }
    except Exception as e:
        import traceback
        traceback.print_exc()
        raise HTTPException(status_code=500, detail=str(e))

@router.get("/presets/{exam_type}")
async def get_exam_preset(exam_type: str):
    """Get default preset for exam type for frontend to display"""
    return rubric_service.get_preset(exam_type)

@router.post("/{rubric_id}/generate-exam")
async def generate_exam(
    rubric_id: str,
    db: Session = Depends(get_db)
):
    """
    Trigger generation of full exam based on rubric configuration.
    Runs asynchronously.
    """
    from ...services.generation_manager import generation_manager
    
    try:
        batch_id = await generation_manager.start_rubric_generation(rubric_id, db)
        return {"batch_id": batch_id, "status": "queued", "message": "Exam generation started"}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

@router.get("/generation-status/{batch_id}")
async def get_generation_status(batch_id: str):
    """Check status of generation batch"""
    from ...services.generation_manager import generation_manager
    return generation_manager.get_status(batch_id)

@router.get("/{rubric_id}/latest-batch")
async def get_latest_batch(rubric_id: str, db: Session = Depends(get_db)):
    """Get the latest generated batch for a rubric"""
    batch = db.query(database.GeneratedBatch).filter(
        database.GeneratedBatch.rubric_id == rubric_id
    ).order_by(database.GeneratedBatch.generated_at.desc()).first()
    
    if not batch:
        raise HTTPException(status_code=404, detail="No generation found for this rubric")
        
    return {"batch_id": batch.id}

@router.get("/{rubric_id}/export-docx")
async def export_rubric_docx(rubric_id: str, db: Session = Depends(get_db)):
    """Export the generated questions for a rubric to a DOCX file"""
    rubric = db.query(database.Rubric).filter(database.Rubric.id == rubric_id).first()
    if not rubric:
        raise HTTPException(status_code=404, detail="Rubric not found")

    batch = db.query(database.GeneratedBatch).filter(
        database.GeneratedBatch.rubric_id == rubric_id
    ).order_by(database.GeneratedBatch.generated_at.desc()).first()
    
    if not batch:
        raise HTTPException(status_code=404, detail="No generated questions found for this rubric")

    questions = db.query(database.Question).filter(
        database.Question.batch_id == batch.id,
        database.Question.status.in_(["valid", "approved", "pending"])
    ).order_by(database.Question.id).all()

    doc = Document()
    doc.add_heading(f"Assessment: {rubric.title}", 0)
    
    for i, q in enumerate(questions):
        # Default question text cleanup
        q_text = str(q.question_text)
        if q_text.strip().startswith('{'):
            try:
                parsed = json.loads(q_text)
                inner = parsed.get("questions", [parsed])[0]
                q_text = inner.get("question_text", inner.get("questionText", q_text))
            except:
                pass
        else:
            import re
            match = re.search(r'"question_text"\s*:\s*"((?:[^"\\]|\\.)*)"', q_text)
            if match:
                q_text = match.group(1).replace('\\"', '"').replace('\\n', '\n')
                
        p = doc.add_paragraph()
        p.add_run(f"Q{i+1}. ").bold = True
        p.add_run(q_text)
        
        # If it's MCQ, add options
        options = []
        if q.options:
            try:
                opts = json.loads(q.options)
                if isinstance(opts, dict):
                    options = [f"{k}) {v}" for k, v in opts.items()]
                elif isinstance(opts, list):
                    options = [str(o) for o in opts]
            except Exception:
                pass
                
        if options:
            for opt in options:
                doc.add_paragraph(opt)
                
        if q.correct_answer:
            p_ans = doc.add_paragraph()
            p_ans.add_run(f"Answer: {q.correct_answer}").italic = True
        
        type_str = str(getattr(q, 'question_type', '') or '').upper()
        diff_str = str(getattr(q, 'difficulty', 'Unknown') or 'Unknown').title()
        bloom_str = str(getattr(q, 'bloom_level', 'Unknown') or 'Unknown').title()
        co_str = getattr(q, 'co_id', 'N/A') or 'N/A'
        lo_str = getattr(q, 'lo_id', 'N/A') or 'N/A'
        
        # Meta properties
        doc.add_paragraph(f"Type: {type_str} | Marks: {getattr(q, 'marks', 1)} | Difficulty: {diff_str} | Bloom: {bloom_str}")
        doc.add_paragraph(f"Mapping: CO: {co_str}, LO: {lo_str}")
        doc.add_paragraph("") # Space

    # Save to BytesIO buffer
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    
    filename = f"{rubric.title.replace(' ', '_')}_Questions.docx"
    
    return Response(
        content=buffer.getvalue(),
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={
            "Content-Disposition": f"attachment; filename={filename}"
        }
    )

@router.get("/{rubric_id}/export-pdf")
async def export_rubric_pdf(rubric_id: str, db: Session = Depends(get_db)):
    """Export the generated questions for a rubric to a PDF file"""
    try:
        from fpdf import FPDF
    except ImportError:
        raise HTTPException(status_code=500, detail="PDF generation library (fpdf2) is not installed")

    rubric = db.query(database.Rubric).filter(database.Rubric.id == rubric_id).first()
    if not rubric:
        raise HTTPException(status_code=404, detail="Rubric not found")

    batch = db.query(database.GeneratedBatch).filter(
        database.GeneratedBatch.rubric_id == rubric_id
    ).order_by(database.GeneratedBatch.generated_at.desc()).first()
    
    if not batch:
        raise HTTPException(status_code=404, detail="No generated questions found for this rubric")

    questions = db.query(database.Question).filter(
        database.Question.batch_id == batch.id,
        database.Question.status.in_(["valid", "approved", "pending"])
    ).order_by(database.Question.id).all()

    class PDF(FPDF):
        def header(self):
            self.set_font("helvetica", "B", 15)
            self.cell(0, 10, f"Assessment: {rubric.title}", align="C")
            self.ln(15)

        def footer(self):
            self.set_y(-15)
            self.set_font("helvetica", "I", 8)
            self.cell(0, 10, f"Page {self.page_no()}", align="C")

    pdf = PDF()
    pdf.add_page()
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.set_font("helvetica", size=11)

    for i, q in enumerate(questions):
        # Default question text cleanup
        q_text = str(q.question_text)
        if q_text.strip().startswith('{'):
            try:
                parsed = json.loads(q_text)
                inner = parsed.get("questions", [parsed])[0]
                q_text = inner.get("question_text", inner.get("questionText", q_text))
            except:
                pass
        else:
            import re
            match = re.search(r'"question_text"\s*:\s*"((?:[^"\\]|\\.)*)"', q_text)
            if match:
                q_text = match.group(1).replace('\\"', '"').replace('\\n', '\n')

        # Clean up stray unicode quotes that FPDF might choke on
        q_text = q_text.replace('\u2019', "'").replace('\u2018', "'").replace('\u201c', '"').replace('\u201d', '"')

        pdf.set_font("helvetica", "B", 11)
        pdf.cell(10, 8, f"Q{i+1}.", ln=0)
        pdf.set_font("helvetica", "", 11)
        pdf.multi_cell(0, 8, q_text)
        
        # Options
        options = []
        if q.options:
            try:
                opts = json.loads(q.options)
                if isinstance(opts, dict):
                    options = [f"{k}) {v}" for k, v in opts.items()]
                elif isinstance(opts, list):
                    options = [str(o) for o in opts]
            except Exception:
                pass
                
        if options:
            pdf.set_font("helvetica", "", 10)
            for opt in options:
                pdf.set_x(25)
                pdf.multi_cell(0, 6, str(opt).replace('\u2019', "'").replace('\u2018', "'").replace('\u201c', '"').replace('\u201d', '"'))
                
        # Answer
        if q.correct_answer:
            pdf.set_font("helvetica", "I", 10)
            pdf.set_x(25)
            pdf.multi_cell(0, 6, f"Answer: {q.correct_answer}")
        
        # Meta info
        pdf.set_font("helvetica", "I", 9)
        pdf.set_text_color(100, 100, 100)
        
        type_str = str(getattr(q, 'question_type', '') or '').upper()
        diff_str = str(getattr(q, 'difficulty', 'Unknown') or 'Unknown').title()
        bloom_str = str(getattr(q, 'bloom_level', 'Unknown') or 'Unknown').title()
        co_str = getattr(q, 'co_id', 'N/A') or 'N/A'
        lo_str = getattr(q, 'lo_id', 'N/A') or 'N/A'
        
        meta = f"Type: {type_str} | Marks: {getattr(q, 'marks', 1)} | Difficulty: {diff_str} | Bloom: {bloom_str}"
        pdf.cell(0, 6, meta, ln=1)
        
        mapping = f"Mapping: CO: {co_str}, LO: {lo_str}"
        pdf.cell(0, 6, mapping, ln=1)
        
        pdf.set_text_color(0, 0, 0)
        pdf.ln(4)

    # Save to BytesIO buffer
    pdf_bytes = pdf.output()
    if isinstance(pdf_bytes, str):
        pdf_bytes = pdf_bytes.encode('latin1')
    
    filename = f"{rubric.title.replace(' ', '_')}_Questions.pdf"
    
    return Response(
        content=bytes(pdf_bytes),
        media_type="application/pdf",
        headers={
            "Content-Disposition": f"attachment; filename={filename}"
        }
    )
